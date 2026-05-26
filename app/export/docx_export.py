from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.db.database import rows_to_dicts


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_font(run, name: str = "宋体", size: int | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for name, size, color in [("Title", 24, "172033"), ("Heading 1", 16, "172033"), ("Heading 2", 13, "2563EB")]:
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)


def _add_body(doc: Document, text: str) -> None:
    for raw in (text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(22)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(4)
        _set_font(p.add_run(line))


def _add_table(doc: Document, headers: list[str], rows: list[list[Any]], title: str | None = None) -> None:
    if title:
        _add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        _set_cell_shading(cell, "EAF2FF")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value or "")
    doc.add_paragraph()


def export_project_docx(conn, project_id: int, output: str | Path, model_name: str = "") -> Path:
    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    project = dict(conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone())
    chapters = rows_to_dicts(
        conn.execute("SELECT * FROM chapters WHERE project_id = ? ORDER BY chapter_number", (project_id,)).fetchall()
    )
    characters = rows_to_dicts(
        conn.execute("SELECT * FROM characters WHERE project_id = ? ORDER BY importance DESC, id", (project_id,)).fetchall()
    )
    world_rules = rows_to_dicts(conn.execute("SELECT * FROM world_rules WHERE project_id = ? ORDER BY id", (project_id,)).fetchall())
    foreshadows = rows_to_dicts(conn.execute("SELECT * FROM foreshadows WHERE project_id = ? ORDER BY id", (project_id,)).fetchall())
    summaries = rows_to_dicts(conn.execute("SELECT * FROM chapter_summaries WHERE project_id = ? ORDER BY chapter_id", (project_id,)).fetchall())
    quality_reports = rows_to_dicts(
        conn.execute("SELECT * FROM quality_reports WHERE project_id = ? ORDER BY id DESC LIMIT 30", (project_id,)).fetchall()
    )

    doc = Document()
    _style_document(doc)

    title = project.get("title") or project.get("name") or "未命名作品"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    _set_font(run, "微软雅黑", 26)
    run.font.color.rgb = RGBColor(23, 32, 51)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(project.get("genre") or "")
    sub_run.italic = True
    _set_font(sub_run, "微软雅黑", 12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}    使用模型：{model_name or '未记录'}"), "微软雅黑", 10)
    doc.add_page_break()

    _add_heading(doc, "小说简介")
    _add_body(doc, project.get("description") or "")

    _add_table(
        doc,
        ["人物", "身份", "性格/动机", "状态"],
        [[c.get("name"), c.get("role"), (c.get("personality") or c.get("motivation") or "")[:120], c.get("status")] for c in characters[:20]],
        "主要人物表",
    )
    _add_table(
        doc,
        ["类型", "规则"],
        [[w.get("category"), w.get("rule_text")] for w in world_rules[:20]],
        "世界观设定",
    )
    if foreshadows:
        _add_table(
            doc,
            ["伏笔", "状态", "预计回收", "风险"],
            [[f.get("name"), f.get("status"), f.get("expected_resolution_chapter"), f.get("risk_note")] for f in foreshadows[:20]],
            "伏笔表",
        )

    _add_heading(doc, "目录")
    for chapter in chapters:
        _set_font(doc.add_paragraph().add_run(f"第 {chapter['chapter_number']} 章  {chapter['title']}"))
    doc.add_page_break()

    for chapter in chapters:
        _add_heading(doc, f"第 {chapter['chapter_number']} 章  {chapter['title']}")
        _add_body(doc, chapter.get("content") or "")
        if chapter != chapters[-1]:
            doc.add_page_break()

    doc.add_section(WD_SECTION.NEW_PAGE)
    _add_heading(doc, "附录：Story Memory 摘要")
    for item in summaries[:20]:
        _add_heading(doc, f"章节摘要 #{item.get('chapter_id')}", 2)
        _add_body(doc, item.get("short_summary") or item.get("detailed_summary") or "")
    if quality_reports:
        _add_heading(doc, "质量检查报告摘要")
        rows = []
        for report in quality_reports:
            payload = json.loads(report.get("report_json") or "{}")
            rows.append([report.get("report_type"), report.get("score"), report.get("risk_level"), (payload.get("summary") or "")[:120]])
        _add_table(doc, ["类型", "分数", "风险", "摘要"], rows)

    doc.save(output_path)
    return output_path

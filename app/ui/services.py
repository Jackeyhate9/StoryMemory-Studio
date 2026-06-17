from __future__ import annotations

import os
import socket
import shutil
import zipfile
import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from dotenv import load_dotenv, set_key
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from app.config import get_settings
from app.checkers.consistency import ConsistencyChecker
from app.context.builder import ContextBuilder
from app.db.database import db_session, get_project, init_db, log_edit, log_generation, row_to_dict, rows_to_dicts
from app.db.models import ContextRequest
from app.generation.chapter import generate_chapter as llm_generate_chapter, rewrite_if_too_similar
from app.llm.client import get_llm
from app.memory.extractor import MemoryExtractor
from app.memory.writer import upsert_extraction
from app.style.profiler import analyze_style
from app.style.similarity_guard import check_similarity
from app.style.style_safety import clean_sample_text, read_style_file
from app.style.style_schema import StyleProfileInput
from app.style.style_store import default_style_profile, export_style_profile, get_style_samples, list_style_profiles, save_style_profile
from app.creative_center import (
    adapt_chapter_for_ip,
    adapt_platform_for_chapter,
    analyze_character_arc_for_project,
    analyze_pacing_for_chapter,
    analyze_platform_fit_for_chapter,
    character_presence_for_project,
    detect_ai_tone_for_chapter,
    detect_character_drift_for_chapter,
    ensure_builtin_platform_profiles,
    plan_payoff_for_foreshadow,
    recommend_payoff_for_project,
    rewrite_ai_tone_for_chapter,
)

ROOT = Path(os.environ.get("STORYMEMORY_RUNTIME_DIR", Path(__file__).parents[2])).expanduser().resolve()
RESOURCE_ROOT = Path(os.environ.get("STORYMEMORY_RESOURCE_DIR", Path(__file__).parents[2])).expanduser().resolve()
ENV_PATH = ROOT / ".env"
BACKUP_DIR = ROOT / "backups"

EDITABLE_MEMORY_TABLES: dict[str, dict[str, Any]] = {
    "characters": {
        "label": "人物卡",
        "fields": [
            "name",
            "role",
            "appearance",
            "personality",
            "motivation",
            "secrets",
            "abilities",
            "status",
            "current_location",
            "hard_constraints",
        ],
        "required": ["name"],
    },
    "world_rules": {
        "label": "世界观规则",
        "fields": ["category", "rule_text", "rigidity", "source", "is_active"],
        "required": ["rule_text"],
    },
    "forbidden_rules": {
        "label": "禁止违背设定",
        "fields": ["rule_text", "category", "severity", "source", "is_active"],
        "required": ["rule_text"],
    },
    "locations": {
        "label": "地点",
        "fields": ["name", "type", "description", "rules"],
        "required": ["name"],
    },
    "items": {
        "label": "道具",
        "fields": ["name", "type", "description", "owner", "location", "status", "constraints"],
        "required": ["name"],
    },
    "foreshadows": {
        "label": "伏笔",
        "fields": ["name", "related_thread", "status", "expected_resolution_chapter", "resolution_method", "risk_note", "evidence"],
        "required": ["name"],
    },
    "unresolved_questions": {
        "label": "未解决问题",
        "fields": ["question", "related_thread", "status", "priority", "notes"],
        "required": ["question"],
    },
}

EDITABLE_MEMORY_TABLES.update(
    {
        "characters": {
            "label": "人物卡",
            "fields": [
                "name",
                "role",
                "appearance",
                "personality",
                "motivation",
                "secrets",
                "abilities",
                "status",
                "current_location",
                "hard_constraints",
                "importance",
            ],
            "required": ["name"],
        },
        "world_rules": {
            "label": "世界观规则",
            "fields": ["category", "rule_text", "rigidity", "source", "is_active"],
            "required": ["rule_text"],
        },
        "forbidden_rules": {
            "label": "禁止违背设定",
            "fields": ["rule_text", "category", "severity", "source", "is_active"],
            "required": ["rule_text"],
        },
        "locations": {
            "label": "地点",
            "fields": ["name", "type", "description", "rules"],
            "required": ["name"],
        },
        "organizations": {
            "label": "势力组织",
            "fields": ["name", "type", "description", "leader", "status"],
            "required": ["name"],
        },
        "items": {
            "label": "道具",
            "fields": ["name", "type", "description", "owner", "location", "status", "constraints"],
            "required": ["name"],
        },
        "abilities": {
            "label": "能力体系",
            "fields": ["name", "owner", "system", "description", "limitations", "cost", "level"],
            "required": ["name"],
        },
        "character_relationships": {
            "label": "人物关系",
            "fields": ["character_a_name", "character_b_name", "relationship_type", "status", "description", "evidence"],
            "required": ["character_a_name", "character_b_name"],
        },
        "foreshadows": {
            "label": "伏笔",
            "fields": [
                "name",
                "related_thread",
                "status",
                "expected_resolution_chapter",
                "resolution_method",
                "last_mentioned_chapter_id",
                "risk_note",
                "evidence",
            ],
            "required": ["name"],
        },
        "timeline_events": {
            "label": "时间线事件",
            "fields": ["story_time", "sort_key", "event_text", "location", "characters_json", "duration", "confidence"],
            "required": ["event_text"],
        },
        "chapter_summaries": {
            "label": "章节摘要",
            "fields": ["chapter_id", "short_summary", "detailed_summary", "key_characters_json", "key_locations_json", "plot_threads_json"],
            "required": ["chapter_id"],
        },
        "chapter_facts": {
            "label": "章节事实",
            "fields": ["chapter_id", "fact_type", "subject", "predicate", "object", "fact_text", "certainty", "source_quote", "is_active"],
            "required": ["chapter_id", "fact_text"],
        },
        "style_profiles": {
            "label": "风格画像",
            "fields": ["name", "style_name", "platform", "pov", "safe_style_summary", "do_rules_json", "dont_rules_json", "profile_json", "is_default"],
            "required": ["name"],
        },
        "unresolved_questions": {
            "label": "未解决问题",
            "fields": ["question", "related_thread", "status", "priority", "notes"],
            "required": ["question"],
        },
    }
)


def bootstrap() -> None:
    load_dotenv(ENV_PATH)
    (ROOT / "data").mkdir(exist_ok=True)
    (ROOT / "exports").mkdir(exist_ok=True)
    (ROOT / "logs").mkdir(exist_ok=True)
    BACKUP_DIR.mkdir(exist_ok=True)
    init_db()


def save_env(values: dict[str, str]) -> None:
    if not ENV_PATH.exists():
        ENV_PATH.write_text("", encoding="utf-8")
    for key, value in values.items():
        set_key(str(ENV_PATH), key, value)
        os.environ[key] = value
    get_settings.cache_clear()


def is_port_open(host: str, port: int, timeout: float = 0.3) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.settimeout(timeout)
        return sock.connect_ex((host, port)) == 0


def health_snapshot() -> dict[str, Any]:
    settings = get_settings()
    db_path = Path(settings.db_path)
    data_dir = db_path.parent.resolve()
    ollama_host = settings.ollama_base_url.replace("http://", "").replace("https://", "").split("/")[0]
    ollama_ok = False
    if ":" in ollama_host:
        host, port_text = ollama_host.rsplit(":", 1)
        try:
            ollama_ok = is_port_open(host, int(port_text))
        except ValueError:
            ollama_ok = False
    return {
        "database_path": str(db_path),
        "database_exists": db_path.exists(),
        "data_dir": str(data_dir),
        "data_dir_exists": data_dir.exists(),
        "database_size_mb": round(db_path.stat().st_size / 1024 / 1024, 2) if db_path.exists() else 0,
        "on_system_drive": str(data_dir).lower().startswith(str(Path.home().drive).lower()) if Path.home().drive else False,
        "env_exists": ENV_PATH.exists(),
        "default_provider": settings.llm_provider,
        "deepseek_configured": bool(settings.deepseek_api_key),
        "glm_configured": bool(settings.glm_api_key or settings.zai_api_key or settings.bigmodel_api_key),
        "openai_configured": bool(settings.openai_api_key or settings.openai_compatible_api_key),
        "ollama_reachable": ollama_ok,
        "local_first": True,
    }


def available_provider_options() -> list[str]:
    settings = get_settings()
    options = ["auto", "none"]
    if settings.deepseek_api_key:
        options.append("deepseek")
    if settings.glm_api_key or settings.zai_api_key or settings.bigmodel_api_key:
        options.append("glm")
    if settings.openai_api_key:
        options.append("openai")
    if settings.openai_compatible_api_key:
        options.append("openai_compatible")
    if health_snapshot()["ollama_reachable"]:
        options.append("ollama")
    for item in ["deepseek", "glm", "openai", "openai_compatible", "ollama"]:
        if item not in options:
            options.append(item)
    return options


def resolve_auto_provider(provider: str) -> str:
    if provider != "auto":
        return provider
    settings = get_settings()
    preferred = settings.llm_provider
    if preferred == "deepseek" and settings.deepseek_api_key:
        return "deepseek"
    if preferred in {"glm", "zhipu", "zai", "bigmodel"} and (
        settings.glm_api_key or settings.zai_api_key or settings.bigmodel_api_key
    ):
        return "glm"
    if preferred == "openai" and settings.openai_api_key:
        return "openai"
    if preferred == "openai_compatible" and settings.openai_compatible_api_key:
        return "openai_compatible"
    if preferred == "ollama" and health_snapshot()["ollama_reachable"]:
        return "ollama"
    if settings.deepseek_api_key:
        return "deepseek"
    if settings.glm_api_key or settings.zai_api_key or settings.bigmodel_api_key:
        return "glm"
    if settings.openai_compatible_api_key:
        return "openai_compatible"
    if settings.openai_api_key:
        return "openai"
    if health_snapshot()["ollama_reachable"]:
        return "ollama"
    return "none"


def suggested_data_dirs() -> list[str]:
    candidates = [
        ROOT / "data",
        Path.home() / "Documents" / "长篇记忆小说",
    ]
    for drive in ["D:", "E:", "F:"]:
        root = Path(f"{drive}/")
        if root.exists():
            candidates.append(root / "长篇记忆小说")
    seen: set[str] = set()
    result: list[str] = []
    for path in candidates:
        text = str(path.resolve())
        if text not in seen:
            seen.add(text)
            result.append(text)
    return result


def migrate_data_storage(new_data_dir: str, copy_existing: bool = True) -> Path:
    target_dir = Path(new_data_dir).expanduser().resolve()
    if not str(target_dir):
        raise ValueError("数据目录不能为空")
    target_dir.mkdir(parents=True, exist_ok=True)
    target_db = target_dir / "storymemory.sqlite3"
    current_db = Path(get_settings().db_path).expanduser().resolve()
    if current_db == target_db:
        init_db(target_db)
        save_env({"STORYMEMORY_DB_PATH": str(target_db)})
        return target_db
    if copy_existing and current_db.exists() and not target_db.exists():
        shutil.copy2(current_db, target_db)
    init_db(target_db)
    save_env({"STORYMEMORY_DB_PATH": str(target_db)})
    return target_db


def list_projects() -> list[dict[str, Any]]:
    init_db()
    with db_session() as conn:
        return rows_to_dicts(conn.execute("SELECT * FROM projects ORDER BY updated_at DESC, id DESC").fetchall())


def dashboard_snapshot(project: str | None) -> dict[str, Any]:
    empty = {
        "chapter_count": 0,
        "character_count": 0,
        "open_foreshadow_count": 0,
        "last_generation": "",
        "last_consistency": "",
        "next_action": "请先从 0 创建小说或导入已有章节。",
    }
    if not project:
        return empty
    with db_session() as conn:
        p = get_project(conn, project)
        chapter_count = conn.execute("SELECT COUNT(*) AS n FROM chapters WHERE project_id = ?", (p["id"],)).fetchone()["n"]
        character_count = conn.execute("SELECT COUNT(*) AS n FROM characters WHERE project_id = ?", (p["id"],)).fetchone()["n"]
        open_foreshadow_count = conn.execute(
            "SELECT COUNT(*) AS n FROM foreshadows WHERE project_id = ? AND status NOT IN ('resolved','已回收','废弃','abandoned')",
            (p["id"],),
        ).fetchone()["n"]
        last_log = conn.execute("SELECT operation, status, created_at FROM generation_logs WHERE project_id = ? ORDER BY id DESC LIMIT 1", (p["id"],)).fetchone()
        last_check = conn.execute(
            "SELECT risk_level, score, created_at FROM quality_reports WHERE project_id = ? AND report_type IN ('consistency','ai_tone_detector','pacing_analyzer') ORDER BY id DESC LIMIT 1",
            (p["id"],),
        ).fetchone()
        facts = conn.execute("SELECT COUNT(*) AS n FROM chapter_facts WHERE project_id = ?", (p["id"],)).fetchone()["n"]
    if chapter_count == 0:
        next_action = "当前项目还没有章节。建议导入第一章，或从 0 创建小说生成第一章。"
    elif facts == 0:
        next_action = "已有章节但记忆事实较少。建议进入「章节导入」或「记忆库看板」检查抽取结果。"
    elif open_foreshadow_count >= 5:
        next_action = "未回收伏笔较多。建议进入「伏笔回收推荐」安排回收窗口。"
    elif last_check and last_check["risk_level"] in {"medium", "high"}:
        next_action = "最近质量检查存在风险。建议进入「AI 腔检测」或「剧情节奏诊断」。"
    else:
        next_action = "项目状态正常。可以继续进入「章节生成」推进下一章。"
    return {
        "chapter_count": chapter_count,
        "character_count": character_count,
        "open_foreshadow_count": open_foreshadow_count,
        "last_generation": f"{last_log['operation']} {last_log['created_at']}" if last_log else "",
        "last_consistency": f"{last_check['risk_level']} / {last_check['score']}" if last_check else "",
        "next_action": next_action,
    }


def create_or_update_project(
    name: str,
    title: str,
    description: str = "",
    genre: str = "",
    platform: str = "",
    default_model: str = "",
) -> None:
    metadata = {"default_model": default_model} if default_model else {}
    with db_session() as conn:
        before = row_to_dict(conn.execute("SELECT * FROM projects WHERE name = ?", (name,)).fetchone())
        conn.execute(
            """
            INSERT INTO projects (name, title, description, genre, target_platform, metadata_json)
            VALUES (?, ?, ?, ?, ?, ?)
            ON CONFLICT(name) DO UPDATE SET
              title=excluded.title,
              description=excluded.description,
              genre=excluded.genre,
              target_platform=excluded.target_platform,
              metadata_json=excluded.metadata_json,
              updated_at=CURRENT_TIMESTAMP
            """,
            (name, title, description, genre, platform, __import__("json").dumps(metadata, ensure_ascii=False)),
        )
        project = get_project(conn, name)
        log_edit(conn, project["id"], "projects", project["id"], "upsert", before, project, "项目管理页保存")


def project_by_name(name: str) -> dict[str, Any] | None:
    with db_session() as conn:
        return row_to_dict(conn.execute("SELECT * FROM projects WHERE name = ?", (name,)).fetchone())


def archive_project(name: str) -> None:
    with db_session() as conn:
        project = get_project(conn, name)
        before = dict(project)
        conn.execute("UPDATE projects SET status = 'archived', updated_at = CURRENT_TIMESTAMP WHERE id = ?", (project["id"],))
        after = row_to_dict(conn.execute("SELECT * FROM projects WHERE id = ?", (project["id"],)).fetchone())
        log_edit(conn, project["id"], "projects", project["id"], "archive", before, after, "归档项目")


def delete_project(name: str) -> None:
    with db_session() as conn:
        project = get_project(conn, name)
        before = dict(project)
        log_edit(conn, project["id"], "projects", project["id"], "delete", before, {}, "删除项目及级联数据")
        conn.execute("DELETE FROM projects WHERE id = ?", (project["id"],))


def project_id_for_name(name: str) -> int:
    with db_session() as conn:
        return int(get_project(conn, name)["id"])


def import_chapter_text(
    project: str,
    number: int,
    title: str,
    content: str,
    volume: str = "",
    provider: str = "none",
) -> dict[str, Any]:
    with db_session() as conn:
        p = get_project(conn, project)
        conn.execute(
            """
            INSERT INTO chapters (project_id, chapter_number, volume, title, content, word_count)
            VALUES (?, ?, ?, ?, ?, ?)
            ON CONFLICT(project_id, chapter_number) DO UPDATE SET
              volume=excluded.volume, title=excluded.title, content=excluded.content,
              word_count=excluded.word_count, updated_at=CURRENT_TIMESTAMP
            """,
            (p["id"], number, volume, title, content, len(content)),
        )
        chapter = conn.execute(
            "SELECT * FROM chapters WHERE project_id = ? AND chapter_number = ?", (p["id"], number)
        ).fetchone()
        provider = resolve_auto_provider(provider)
        client = None if provider == "none" else get_llm(provider)
        extraction = MemoryExtractor(client).extract(title, content)
        upsert_extraction(conn, p["id"], chapter["id"], extraction)
        log_generation(
            conn,
            p["id"],
            "ui_import_chapter_extract_memory",
            provider=provider,
            response=extraction.model_dump_json(),
            structured=extraction.model_dump(),
            chapter_id=chapter["id"],
        )
        return extraction.model_dump()


def read_uploaded_text(filename: str, data: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix in {"txt", "md"}:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="ignore")
    if suffix == "docx":
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("当前仅支持 txt、md、docx 文本导入")


def split_chapters_from_text(text: str) -> list[dict[str, Any]]:
    pattern = re.compile(r"(?m)^\s*(第[一二三四五六七八九十百千万零〇\d]+章[^\n\r]*|Chapter\s+\d+[^\n\r]*)\s*$")
    matches = list(pattern.finditer(text))
    if not matches:
        return [{"title": "导入章节", "content": text.strip()}] if text.strip() else []
    chapters: list[dict[str, Any]] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        content = text[start:end].strip()
        if content:
            chapters.append({"title": match.group(1).strip(), "content": content})
    return chapters


def import_chapter_batch(
    project: str,
    start_number: int,
    chapters: list[dict[str, Any]],
    volume: str = "",
    provider: str = "none",
) -> list[dict[str, Any]]:
    results = []
    for offset, chapter in enumerate(chapters):
        number = start_number + offset
        title = chapter.get("title") or f"第 {number} 章"
        content = chapter.get("content") or ""
        if not content.strip():
            continue
        extraction = import_chapter_text(project, number, title, content, volume, provider)
        results.append({"chapter_number": number, "title": title, "extraction": extraction})
    return results


def chapter_options(project: str) -> list[dict[str, Any]]:
    with db_session() as conn:
        p = get_project(conn, project)
        return rows_to_dicts(
            conn.execute(
                "SELECT id, chapter_number, title, content FROM chapters WHERE project_id = ? ORDER BY chapter_number",
                (p["id"],),
            ).fetchall()
        )


def character_options(project: str) -> list[dict[str, Any]]:
    with db_session() as conn:
        p = get_project(conn, project)
        return rows_to_dicts(conn.execute("SELECT id, name, role, status FROM characters WHERE project_id = ? ORDER BY importance DESC, id", (p["id"],)).fetchall())


def foreshadow_options(project: str) -> list[dict[str, Any]]:
    with db_session() as conn:
        p = get_project(conn, project)
        return rows_to_dicts(conn.execute("SELECT id, name, status, risk_note FROM foreshadows WHERE project_id = ? ORDER BY status, id", (p["id"],)).fetchall())


def run_ai_tone(project: str, chapter_id: int) -> dict[str, Any]:
    return detect_ai_tone_for_chapter(project_id_for_name(project), chapter_id)


def run_ai_tone_rewrite(project: str, chapter_id: int, apply: bool = False) -> dict[str, Any]:
    return rewrite_ai_tone_for_chapter(project_id_for_name(project), chapter_id, apply)


def run_pacing(project: str, chapter_id: int) -> dict[str, Any]:
    return analyze_pacing_for_chapter(project_id_for_name(project), chapter_id)


def run_payoff_recommend(project: str) -> dict[str, Any]:
    return recommend_payoff_for_project(project_id_for_name(project))


def run_payoff_plan(project: str, foreshadow_id: int) -> dict[str, Any]:
    return plan_payoff_for_foreshadow(project_id_for_name(project), foreshadow_id)


def run_character_arc(project: str, character_id: int) -> dict[str, Any]:
    return analyze_character_arc_for_project(project_id_for_name(project), character_id)


def run_character_drift(project: str, chapter_id: int) -> dict[str, Any]:
    return detect_character_drift_for_chapter(project_id_for_name(project), chapter_id)


def run_character_presence(project: str) -> list[dict[str, Any]]:
    return character_presence_for_project(project_id_for_name(project))


def run_platform_fit(project: str, chapter_id: int, platform: str) -> dict[str, Any]:
    ensure_builtin_platform_profiles()
    return analyze_platform_fit_for_chapter(project_id_for_name(project), chapter_id, platform)


def run_platform_adapt(project: str, chapter_id: int, platform: str, apply: bool = False) -> dict[str, Any]:
    return adapt_platform_for_chapter(project_id_for_name(project), chapter_id, platform, apply)


def run_adaptation_matrix(project: str, chapter_id: int, adaptation_type: str = "all") -> dict[str, Any]:
    return adapt_chapter_for_ip(project_id_for_name(project), chapter_id, adaptation_type)


def get_chapter(project: str, chapter_id: int) -> dict[str, Any] | None:
    with db_session() as conn:
        p = get_project(conn, project)
        return row_to_dict(conn.execute("SELECT * FROM chapters WHERE project_id = ? AND id = ?", (p["id"], chapter_id)).fetchone())


def update_chapter(project: str, chapter_id: int, title: str, content: str, volume: str = "", outline: str = "") -> None:
    with db_session() as conn:
        p = get_project(conn, project)
        before = row_to_dict(conn.execute("SELECT * FROM chapters WHERE project_id = ? AND id = ?", (p["id"], chapter_id)).fetchone())
        conn.execute(
            """
            UPDATE chapters
            SET title = ?, content = ?, volume = ?, outline = ?, word_count = ?, updated_at = CURRENT_TIMESTAMP
            WHERE project_id = ? AND id = ?
            """,
            (title, content, volume, outline, len(content), p["id"], chapter_id),
        )
        after = row_to_dict(conn.execute("SELECT * FROM chapters WHERE project_id = ? AND id = ?", (p["id"], chapter_id)).fetchone())
        log_edit(conn, p["id"], "chapters", chapter_id, "update", before, after, "章节编辑页保存")


def delete_chapter(project: str, chapter_id: int) -> None:
    with db_session() as conn:
        p = get_project(conn, project)
        before = row_to_dict(conn.execute("SELECT * FROM chapters WHERE project_id = ? AND id = ?", (p["id"], chapter_id)).fetchone())
        log_edit(conn, p["id"], "chapters", chapter_id, "delete", before, {}, "删除章节")
        conn.execute("DELETE FROM chapters WHERE project_id = ? AND id = ?", (p["id"], chapter_id))


def export_all_chapters_markdown(project: str) -> tuple[bytes, str, str]:
    chapters = chapter_options(project)
    parts = []
    for chapter in chapters:
        parts.append(f"# 第 {chapter['chapter_number']} 章 {chapter['title']}\n\n{chapter.get('content') or ''}")
    filename = f"{project}_chapters.md"
    return "\n\n---\n\n".join(parts).encode("utf-8-sig"), filename, "text/markdown"


def export_all_chapters_json(project: str) -> tuple[bytes, str, str]:
    chapters = chapter_options(project)
    filename = f"{project}_chapters.json"
    return json.dumps(chapters, ensure_ascii=False, indent=2).encode("utf-8"), filename, "application/json"


def _chapter_filename(chapter_number: int, title: str, suffix: str) -> str:
    safe = "".join(ch for ch in title if ch.isalnum() or ch in "-_（）()[]【】 ").strip()[:60] or "chapter"
    return f"{chapter_number:04d}_{safe}.{suffix}"


def export_chapter_bytes(title: str, content: str, chapter_number: int, file_format: str) -> tuple[bytes, str, str]:
    if file_format == "txt":
        filename = _chapter_filename(chapter_number, title, "txt")
        return f"{title}\n\n{content}".encode("utf-8-sig"), filename, "text/plain"
    if file_format == "doc":
        filename = _chapter_filename(chapter_number, title, "doc")
        escaped_title = title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        paragraphs = "\n".join(
            f"<p>{line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')}</p>"
            for line in content.splitlines()
            if line.strip()
        )
        html = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{escaped_title}</title>
  <style>
    body {{ font-family: SimSun, Microsoft YaHei, serif; font-size: 12pt; line-height: 1.8; }}
    h1 {{ font-family: Microsoft YaHei, SimHei, sans-serif; font-size: 20pt; }}
    p {{ text-indent: 2em; margin: 0 0 8pt 0; }}
  </style>
</head>
<body>
  <h1>{escaped_title}</h1>
  {paragraphs}
</body>
</html>"""
        return html.encode("utf-8-sig"), filename, "application/msword"
    if file_format == "docx":
        filename = _chapter_filename(chapter_number, title, "docx")
        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "宋体"
        styles["Normal"].font.size = Pt(12)
        heading = document.add_heading(title, level=1)
        heading.style.font.name = "微软雅黑"
        for raw in content.splitlines():
            line = raw.strip()
            if not line:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.add_run(line)
        output = BytesIO()
        document.save(output)
        return output.getvalue(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if file_format == "md":
        filename = _chapter_filename(chapter_number, title, "md")
        return f"# {title}\n\n{content}".encode("utf-8-sig"), filename, "text/markdown"
    if file_format == "json":
        filename = _chapter_filename(chapter_number, title, "json")
        payload = {"chapter_number": chapter_number, "title": title, "content": content}
        return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"), filename, "application/json"
    raise ValueError(f"Unsupported export format: {file_format}")


def memory_dashboard(project: str, search: str = "") -> dict[str, list[dict[str, Any]]]:
    tables = {
        "characters": "name || ' ' || role || ' ' || personality || ' ' || status",
        "character_relationships": "character_a_name || ' ' || character_b_name || ' ' || relationship_type || ' ' || description",
        "world_rules": "category || ' ' || rule_text",
        "locations": "name || ' ' || type || ' ' || description",
        "organizations": "name || ' ' || type || ' ' || description || ' ' || status",
        "abilities": "name || ' ' || owner || ' ' || system || ' ' || description",
        "items": "name || ' ' || type || ' ' || owner || ' ' || status",
        "foreshadows": "name || ' ' || status || ' ' || related_thread || ' ' || risk_note",
        "timeline_events": "story_time || ' ' || event_text || ' ' || location",
        "unresolved_questions": "question || ' ' || related_thread || ' ' || status",
        "chapter_summaries": "short_summary || ' ' || detailed_summary",
        "chapter_facts": "fact_text || ' ' || subject || ' ' || object",
    }
    result = {}
    with db_session() as conn:
        p = get_project(conn, project)
        for table, expr in tables.items():
            if search:
                rows = conn.execute(
                    f"SELECT * FROM {table} WHERE project_id = ? AND ({expr}) LIKE ? ORDER BY id DESC LIMIT 200",
                    (p["id"], f"%{search}%"),
                ).fetchall()
            else:
                rows = conn.execute(f"SELECT * FROM {table} WHERE project_id = ? ORDER BY id DESC LIMIT 200", (p["id"],)).fetchall()
            result[table] = rows_to_dicts(rows)
    return result


def build_context(project: str, chapter_number: int, goal: str, outline: str, characters: list[str], locations: list[str], mode: str) -> str:
    with db_session() as conn:
        p = get_project(conn, project)
        req = ContextRequest(
            project=project,
            chapter_number=chapter_number,
            chapter_goal=goal,
            chapter_outline=outline,
            characters=characters,
            locations=locations,
            mode=mode,  # type: ignore[arg-type]
        )
        return ContextBuilder(conn, p).build(req)


def generate_with_context(provider: str, context: str, mode: str, extra_instruction: str = "", project: str | None = None) -> str:
    provider = resolve_auto_provider(provider)
    client = get_llm(provider if provider != "default" else None)
    result = llm_generate_chapter(client, context, mode=mode, extra_instruction=extra_instruction)
    if project:
        style = default_style_profile(project)
        if style:
            samples = get_style_samples(int(style["id"]))
            sample_text = "\n\n".join(
                (row.get("sample_text") or row.get("sample_excerpt_safe") or "")
                for row in samples
            ).strip()
            if sample_text:
                result, report = rewrite_if_too_similar(
                    client,
                    result,
                    sample_text,
                    context,
                    style.get("safe_style_summary") or style.get("profile_json") or "",
                    extra_instruction,
                )
                if report.get("rewrite_required"):
                    with db_session() as conn:
                        p = get_project(conn, project)
                        log_generation(
                            conn,
                            p["id"],
                            "style_similarity_guard_rewrite",
                            provider=provider,
                            prompt=context,
                            response=result,
                            structured=report,
                        )
    return result


def analyze_style_for_project(
    project: str,
    style_name: str,
    samples: list[str],
    target_usage: list[str],
    source_note: str = "",
    save_source: bool = False,
    set_default: bool = False,
    provider: str = "none",
) -> dict[str, Any]:
    provider = resolve_auto_provider(provider)
    input_data = StyleProfileInput(
        style_name=style_name,
        samples=[clean_sample_text(x) for x in samples if clean_sample_text(x)],
        target_usage=target_usage,
        source_note=source_note,
        save_source=save_source,
        set_default=set_default,
    )
    result = analyze_style(input_data, provider)
    return {"input": input_data.model_dump(), "profile": result.model_dump()}


def save_style_analysis(project: str, analysis: dict[str, Any], save_source: bool, set_default: bool) -> int:
    from app.style.style_schema import StyleProfileResult

    profile = StyleProfileResult.model_validate(analysis["profile"])
    samples = analysis.get("input", {}).get("samples", [])
    source_note = analysis.get("input", {}).get("source_note", "")
    return save_style_profile(project, profile, samples, source_note, save_source, set_default)


def style_profiles_for_project(project: str) -> list[dict[str, Any]]:
    return list_style_profiles(project)


def check_style_similarity(sample_text: str, generated_text: str) -> dict[str, Any]:
    return check_similarity(sample_text, generated_text).model_dump()


def save_generated_chapter(project: str, number: int, title: str, content: str, volume: str = "", provider: str = "none") -> dict[str, Any]:
    return import_chapter_text(project, number, title, content, volume, provider)


def check_text(project: str, chapter_number: int, text: str, goal: str = "", provider: str = "none") -> dict[str, Any]:
    provider = resolve_auto_provider(provider)
    with db_session() as conn:
        p = get_project(conn, project)
        context = ContextBuilder(conn, p).build(ContextRequest(project=project, chapter_number=chapter_number, chapter_goal=goal))
        client = None if provider == "none" else get_llm(provider)
        return ConsistencyChecker(client).check(context, text, goal).model_dump()


def update_foreshadow(project: str, row_id: int, status: str, expected: int | None, resolution: str, risk_note: str) -> None:
    with db_session() as conn:
        p = get_project(conn, project)
        before = row_to_dict(conn.execute("SELECT * FROM foreshadows WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
        conn.execute(
            """
            UPDATE foreshadows
            SET status = ?, expected_resolution_chapter = ?, resolution_method = ?, risk_note = ?
            WHERE project_id = ? AND id = ?
            """,
            (status, expected, resolution, risk_note, p["id"], row_id),
        )
        after = row_to_dict(conn.execute("SELECT * FROM foreshadows WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
        log_edit(conn, p["id"], "foreshadows", row_id, "update", before, after, "伏笔管理页更新")


def timeline_events(project: str, order_by: str = "chapter") -> list[dict[str, Any]]:
    order = "c.chapter_number, t.id" if order_by == "chapter" else "t.sort_key, t.story_time, t.id"
    with db_session() as conn:
        p = get_project(conn, project)
        rows = conn.execute(
            f"""
            SELECT t.*, c.chapter_number, c.title AS chapter_title
            FROM timeline_events t
            LEFT JOIN chapters c ON c.id = t.chapter_id
            WHERE t.project_id = ?
            ORDER BY {order}
            """,
            (p["id"],),
        ).fetchall()
        return rows_to_dicts(rows)


def upsert_timeline_event(project: str, values: dict[str, Any], row_id: int | None = None) -> int:
    payload = {
        "chapter_id": int(values.get("chapter_id") or 0) or None,
        "story_time": values.get("story_time", ""),
        "sort_key": values.get("sort_key", ""),
        "event_text": values.get("event_text", ""),
        "location": values.get("location", ""),
        "characters_json": values.get("characters_json", "[]"),
        "duration": values.get("duration", ""),
        "confidence": float(values.get("confidence") or 1.0),
    }
    if not payload["event_text"].strip():
        raise ValueError("时间线事件不能为空")
    with db_session() as conn:
        p = get_project(conn, project)
        if row_id:
            before = row_to_dict(conn.execute("SELECT * FROM timeline_events WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
            conn.execute(
                """
                UPDATE timeline_events
                SET chapter_id=?, story_time=?, sort_key=?, event_text=?, location=?, characters_json=?, duration=?, confidence=?
                WHERE project_id=? AND id=?
                """,
                (
                    payload["chapter_id"],
                    payload["story_time"],
                    payload["sort_key"],
                    payload["event_text"],
                    payload["location"],
                    payload["characters_json"],
                    payload["duration"],
                    payload["confidence"],
                    p["id"],
                    row_id,
                ),
            )
            after = row_to_dict(conn.execute("SELECT * FROM timeline_events WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
            log_edit(conn, p["id"], "timeline_events", row_id, "update", before, after, "时间线管理页保存")
            return row_id
        cur = conn.execute(
            """
            INSERT INTO timeline_events
            (project_id, chapter_id, story_time, sort_key, event_text, location, characters_json, duration, confidence)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                p["id"],
                payload["chapter_id"],
                payload["story_time"],
                payload["sort_key"],
                payload["event_text"],
                payload["location"],
                payload["characters_json"],
                payload["duration"],
                payload["confidence"],
            ),
        )
        new_id = int(cur.lastrowid)
        after = row_to_dict(conn.execute("SELECT * FROM timeline_events WHERE project_id = ? AND id = ?", (p["id"], new_id)).fetchone())
        log_edit(conn, p["id"], "timeline_events", new_id, "insert", {}, after, "时间线管理页新增")
        return new_id


def delete_timeline_event(project: str, row_id: int) -> None:
    with db_session() as conn:
        p = get_project(conn, project)
        before = row_to_dict(conn.execute("SELECT * FROM timeline_events WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
        log_edit(conn, p["id"], "timeline_events", row_id, "delete", before, {}, "时间线管理页删除")
        conn.execute("DELETE FROM timeline_events WHERE project_id = ? AND id = ?", (p["id"], row_id))


def editable_table_labels() -> dict[str, str]:
    return {key: value["label"] for key, value in EDITABLE_MEMORY_TABLES.items()}


def editable_fields(table: str) -> list[str]:
    if table not in EDITABLE_MEMORY_TABLES:
        raise ValueError(f"Unsupported editable table: {table}")
    return list(EDITABLE_MEMORY_TABLES[table]["fields"])


def list_memory_rows(project: str, table: str, search: str = "") -> list[dict[str, Any]]:
    if table not in EDITABLE_MEMORY_TABLES:
        raise ValueError(f"Unsupported editable table: {table}")
    fields = editable_fields(table)
    expr = " || ' ' || ".join(fields)
    with db_session() as conn:
        p = get_project(conn, project)
        if search:
            rows = conn.execute(
                f"SELECT * FROM {table} WHERE project_id = ? AND ({expr}) LIKE ? ORDER BY id DESC LIMIT 300",
                (p["id"], f"%{search}%"),
            ).fetchall()
        else:
            rows = conn.execute(f"SELECT * FROM {table} WHERE project_id = ? ORDER BY id DESC LIMIT 300", (p["id"],)).fetchall()
        return rows_to_dicts(rows)


def get_memory_row(project: str, table: str, row_id: int) -> dict[str, Any] | None:
    if table not in EDITABLE_MEMORY_TABLES:
        raise ValueError(f"Unsupported editable table: {table}")
    with db_session() as conn:
        p = get_project(conn, project)
        return row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())


def upsert_memory_row(project: str, table: str, values: dict[str, Any], row_id: int | None = None) -> int:
    if table not in EDITABLE_MEMORY_TABLES:
        raise ValueError(f"Unsupported editable table: {table}")
    fields = editable_fields(table)
    required = EDITABLE_MEMORY_TABLES[table].get("required", [])
    payload = {field: values.get(field, "") for field in fields}
    numeric_fields = {
        "chapter_id",
        "is_active",
        "is_default",
        "expected_resolution_chapter",
        "last_mentioned_chapter_id",
        "importance",
    }
    float_fields = {"certainty", "confidence"}
    for field in numeric_fields:
        if field in payload:
            payload[field] = int(payload[field] or 0)
    for field in float_fields:
        if field in payload:
            payload[field] = float(payload[field] or 0)
    for field in required:
        if not str(payload.get(field, "")).strip():
            raise ValueError(f"字段不能为空：{field}")
    with db_session() as conn:
        p = get_project(conn, project)
        if row_id:
            before = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
            assignments = ", ".join(f"{field} = ?" for field in fields)
            conn.execute(
                f"UPDATE {table} SET {assignments} WHERE project_id = ? AND id = ?",
                (*[payload[field] for field in fields], p["id"], row_id),
            )
            after = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
            log_edit(conn, p["id"], table, row_id, "update", before, after, "记忆编辑器保存")
            return row_id
        columns = ["project_id", *fields]
        placeholders = ", ".join("?" for _ in columns)
        cur = conn.execute(
            f"INSERT INTO {table} ({', '.join(columns)}) VALUES ({placeholders})",
            (p["id"], *[payload[field] for field in fields]),
        )
        row_id = int(cur.lastrowid)
        after = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
        log_edit(conn, p["id"], table, row_id, "insert", {}, after, "记忆编辑器新增")
        return row_id


def delete_memory_row(project: str, table: str, row_id: int) -> None:
    if table not in EDITABLE_MEMORY_TABLES:
        raise ValueError(f"Unsupported editable table: {table}")
    with db_session() as conn:
        p = get_project(conn, project)
        before = row_to_dict(conn.execute(f"SELECT * FROM {table} WHERE project_id = ? AND id = ?", (p["id"], row_id)).fetchone())
        log_edit(conn, p["id"], table, row_id, "delete", before, {}, "记忆编辑器删除")
        conn.execute(f"DELETE FROM {table} WHERE project_id = ? AND id = ?", (p["id"], row_id))


def create_backup(note: str = "") -> Path:
    bootstrap()
    settings = get_settings()
    db_path = Path(settings.db_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_note = "".join(ch for ch in note.strip() if ch.isalnum() or ch in "-_")[:32]
    name = f"storymemory_backup_{timestamp}{'_' + safe_note if safe_note else ''}.zip"
    output = BACKUP_DIR / name
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        if db_path.exists():
            zf.write(db_path, "data/storymemory.sqlite3")
        if ENV_PATH.exists():
            zf.write(ENV_PATH, ".env")
        for path in (ROOT / "data").glob("*.json"):
            zf.write(path, f"data/{path.name}")
        for path in (ROOT / "data").glob("*.md"):
            zf.write(path, f"data/{path.name}")
        zf.writestr(
            "backup_manifest.txt",
            f"长篇记忆小说备份\ncreated_at={datetime.now().isoformat(timespec='seconds')}\nnote={note}\n",
        )
    return output


def list_backups() -> list[dict[str, Any]]:
    BACKUP_DIR.mkdir(exist_ok=True)
    rows = []
    for path in sorted(BACKUP_DIR.glob("*.zip"), key=lambda p: p.stat().st_mtime, reverse=True):
        rows.append({"name": path.name, "path": str(path), "size_kb": round(path.stat().st_size / 1024, 1)})
    return rows


def restore_backup(zip_bytes: bytes) -> Path:
    bootstrap()
    before = create_backup("before_restore")
    settings = get_settings()
    db_path = Path(settings.db_path)
    tmp = BACKUP_DIR / f"restore_upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    tmp.write_bytes(zip_bytes)
    with zipfile.ZipFile(tmp, "r") as zf:
        members = set(zf.namelist())
        if "data/storymemory.sqlite3" not in members:
            raise ValueError("备份文件缺少 data/storymemory.sqlite3，无法恢复。")
        db_path.parent.mkdir(parents=True, exist_ok=True)
        extracted = BACKUP_DIR / "restore_storymemory.sqlite3"
        with zf.open("data/storymemory.sqlite3") as src, extracted.open("wb") as dst:
            shutil.copyfileobj(src, dst)
        shutil.copy2(extracted, db_path)
        extracted.unlink(missing_ok=True)
        if ".env" in members:
            with zf.open(".env") as src, ENV_PATH.open("wb") as dst:
                shutil.copyfileobj(src, dst)
    tmp.unlink(missing_ok=True)
    init_db(db_path)
    return before
